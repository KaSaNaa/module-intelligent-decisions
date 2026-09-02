import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# ---------------------------------------------------------------------------
# COLOR PALETTE & STYLING CONSTANTS
# ---------------------------------------------------------------------------
NAVY_HEX = "1B365D"      # Primary Header Accent
SLATE_HEX = "2B4C7E"     # H2 Accent
TEAL_HEX = "008080"      # H3 Accent
CHARCOAL_HEX = "222222"  # Body Text
LIGHT_BG_HEX = "F4F6F9"  # Table / Callout Background
BORDER_HEX = "CBD5E1"    # Table / Code Border
HEADER_BG_HEX = "1B365D"  # Table Header Fill

COLOR_NAVY = RGBColor(0x1B, 0x36, 0x5D)
COLOR_SLATE = RGBColor(0x2B, 0x4C, 0x7E)
COLOR_TEAL = RGBColor(0x00, 0x80, 0x80)
COLOR_CHARCOAL = RGBColor(0x22, 0x22, 0x22)
COLOR_MUTED = RGBColor(0x64, 0x74, 0x8B)

# Helper function to set cell background color
def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Helper function to set cell margins (padding)
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

# Helper function to set cell borders
def set_table_borders(table, hex_color=BORDER_HEX):
    tblPr = table._tbl.tblPr
    borders_elm = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
        f'<w:insideV w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders_elm)

# Helper function for adding styled callout / code boxes
def add_code_box(doc, title, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Set left border thick navy, others none
    tcPr = cell._tc.get_or_add_tcPr()
    borders_elm = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:color="{BORDER_HEX}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:color="{BORDER_HEX}"/>'
        f'<w:left w:val="single" w:sz="24" w:color="{NAVY_HEX}"/>'
        f'<w:right w:val="single" w:sz="4" w:color="{BORDER_HEX}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders_elm)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"/* {title} */\n")
    run_t.font.name = "Consolas"
    run_t.font.size = Pt(9)
    run_t.font.bold = True
    run_t.font.color.rgb = COLOR_NAVY
    
    run_c = p.add_run(code_text)
    run_c.font.name = "Consolas"
    run_c.font.size = Pt(8.5)
    run_c.font.color.rgb = COLOR_CHARCOAL
    
    # Empty paragraph after table for spacing
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(6)

def add_heading_1(doc, text):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    for r in h.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = COLOR_NAVY
    return h

def add_heading_2(doc, text):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    for r in h.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = COLOR_SLATE
    return h

def add_heading_3(doc, text):
    h = doc.add_heading(text, level=3)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(2)
    h.paragraph_format.keep_with_next = True
    for r in h.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = COLOR_TEAL
    return h

def add_body(doc, text, bold_prefix=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Calibri"
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_CHARCOAL
    r_body = p.add_run(text)
    r_body.font.name = "Calibri"
    r_body.font.size = Pt(11)
    r_body.font.color.rgb = COLOR_CHARCOAL
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Calibri"
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_CHARCOAL
    r_body = p.add_run(text)
    r_body.font.name = "Calibri"
    r_body.font.size = Pt(11)
    r_body.font.color.rgb = COLOR_CHARCOAL
    return p

def add_figure(doc, image_path, caption_text):
    if os.path.exists(image_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        run = p_img.add_run()
        run.add_picture(image_path, width=Inches(6.0))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(12)
        p_cap.paragraph_format.keep_with_next = True
        run_cap = p_cap.add_run(caption_text)
        run_cap.font.name = "Calibri"
        run_cap.font.size = Pt(9.5)
        run_cap.font.italic = True
        run_cap.font.color.rgb = COLOR_MUTED

def style_table_headers_and_rows(table, col_widths, alignments):
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header Row
    hdr_cells = table.rows[0].cells
    for i, cell in enumerate(hdr_cells):
        cell.width = Inches(col_widths[i])
        set_cell_background(cell, HEADER_BG_HEX)
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        for p in cell.paragraphs:
            p.alignment = alignments[i]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.font.name = "Calibri"
                r.font.size = Pt(10)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data Rows
    for r_idx, row in enumerate(table.rows[1:], start=1):
        bg_color = LIGHT_BG_HEX if r_idx % 2 == 1 else "FFFFFF"
        for i, cell in enumerate(row.cells):
            cell.width = Inches(col_widths[i])
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            for p in cell.paragraphs:
                p.alignment = alignments[i]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = COLOR_CHARCOAL

print("Helper functions configured successfully.")
